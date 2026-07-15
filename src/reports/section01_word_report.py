"""
Digital Electrical Engineer (DEE)

SECTION 01 WORD REPORT GENERATOR

Section 1:
SELECTION OF MAIN DIMENSIONS

Kopylov asynchronous motor design method.

Document formatting:
    Font              : Times New Roman
    Font size         : 14 pt
    Line spacing      : 1.5
    First line indent : 1.25 cm
    Text alignment    : justified

Equations are generated as native Microsoft Word
Office Math objects using OMML.
"""

from pathlib import Path

from docx import Document

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
)

from docx.shared import (
    Pt,
    Cm,
)

from docx.oxml import OxmlElement

from docx.oxml.ns import (
    qn,
)


# ================================================================
# REPORT SETTINGS
# ================================================================

REPORT_DIRECTORY = Path(
    "reports"
)

REPORT_FILENAME = (
    "Section_01_Main_Dimensions.docx"
)

FONT_NAME = "Times New Roman"

FONT_SIZE_PT = 14

LINE_SPACING = 1.5

FIRST_LINE_INDENT_CM = 1.25


# ================================================================
# TEXT FORMATTING
# ================================================================

def set_run_font(
    run,
    font_name=FONT_NAME,
    font_size=FONT_SIZE_PT,
    bold=False,
    italic=False,
):
    """
    Apply font settings to a Word run.
    """

    run.font.name = font_name

    run.font.size = Pt(
        font_size
    )

    run.font.bold = bold

    run.font.italic = italic

    run_properties = (
        run._element.get_or_add_rPr()
    )

    run_fonts = run_properties.rFonts

    if run_fonts is None:

        run_fonts = OxmlElement(
            "w:rFonts"
        )

        run_properties.insert(
            0,
            run_fonts,
        )

    run_fonts.set(
        qn("w:ascii"),
        font_name,
    )

    run_fonts.set(
        qn("w:hAnsi"),
        font_name,
    )

    run_fonts.set(
        qn("w:eastAsia"),
        font_name,
    )

    run_fonts.set(
        qn("w:cs"),
        font_name,
    )


def format_paragraph(
    paragraph,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=True,
):
    """
    Apply standard report paragraph formatting.
    """

    paragraph.alignment = alignment

    paragraph_format = (
        paragraph.paragraph_format
    )

    paragraph_format.line_spacing = (
        LINE_SPACING
    )

    paragraph_format.space_before = Pt(
        0
    )

    paragraph_format.space_after = Pt(
        0
    )

    if first_line_indent:

        paragraph_format.first_line_indent = Cm(
            FIRST_LINE_INDENT_CM
        )

    else:

        paragraph_format.first_line_indent = Cm(
            0
        )


def add_text_paragraph(
    document,
    text,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=True,
):
    """
    Add standard formatted text paragraph.
    """

    paragraph = document.add_paragraph()

    format_paragraph(
        paragraph=paragraph,
        alignment=alignment,
        first_line_indent=first_line_indent,
    )

    run = paragraph.add_run(
        text
    )

    set_run_font(
        run
    )

    return paragraph


# ================================================================
# SECTION TITLE
# ================================================================

def add_section_title(
    document,
    text,
):
    """
    Add section title.
    """

    paragraph = document.add_paragraph()

    format_paragraph(
        paragraph=paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )

    run = paragraph.add_run(
        text
    )

    set_run_font(
        run=run,
        bold=True,
    )

    return paragraph


# ================================================================
# OMML BASIC ELEMENTS
# ================================================================

def create_math_run(
    text,
):
    """
    Create a native OMML mathematical run.
    """

    math_run = OxmlElement(
        "m:r"
    )

    math_run_properties = OxmlElement(
        "m:rPr"
    )

    math_style = OxmlElement(
        "m:sty"
    )

    math_style.set(
        qn("m:val"),
        "p",
    )

    math_run_properties.append(
        math_style
    )

    math_run.append(
        math_run_properties
    )

    math_text = OxmlElement(
        "m:t"
    )

    math_text.text = str(
        text
    )

    math_run.append(
        math_text
    )

    return math_run


def create_math_text(
    text,
):
    """
    Create simple OMML mathematical text.
    """

    return create_math_run(
        text
    )


# ================================================================
# OMML CONTENT APPENDER
# ================================================================

def append_math_content(
    parent,
    content,
):
    """
    Append mathematical content to an OMML element.

    Supported content:
        str
        OMML element
        list / tuple of str and OMML elements
    """

    if isinstance(
        content,
        str,
    ):

        parent.append(
            create_math_text(
                content
            )
        )

        return


    if isinstance(
        content,
        (list, tuple),
    ):

        for element in content:

            append_math_content(
                parent=parent,
                content=element,
            )

        return


    parent.append(
        content
    )


# ================================================================
# OMML SUBSCRIPT
# ================================================================

def create_subscript(
    base,
    subscript,
):
    """
    Create native OMML lower index.

    Examples:
        P_2
        U_1
        D_a
        K_D
        K_E
        B_delta
        alpha_delta
        k_B
        k_ob1
        n_1
        l_delta
    """

    subscript_element = OxmlElement(
        "m:sSub"
    )

    subscript_properties = OxmlElement(
        "m:sSubPr"
    )

    subscript_element.append(
        subscript_properties
    )


    # BASE

    base_element = OxmlElement(
        "m:e"
    )

    append_math_content(
        parent=base_element,
        content=base,
    )

    subscript_element.append(
        base_element
    )


    # SUBSCRIPT

    subscript_value = OxmlElement(
        "m:sub"
    )

    append_math_content(
        parent=subscript_value,
        content=subscript,
    )

    subscript_element.append(
        subscript_value
    )

    return subscript_element


# ================================================================
# OMML SUPERSCRIPT
# ================================================================

def create_superscript(
    base,
    exponent,
):
    """
    Create native OMML superscript.

    Example:
        D^2
    """

    superscript = OxmlElement(
        "m:sSup"
    )

    superscript_properties = OxmlElement(
        "m:sSupPr"
    )

    superscript.append(
        superscript_properties
    )


    # BASE

    base_element = OxmlElement(
        "m:e"
    )

    append_math_content(
        parent=base_element,
        content=base,
    )

    superscript.append(
        base_element
    )


    # EXPONENT

    exponent_element = OxmlElement(
        "m:sup"
    )

    append_math_content(
        parent=exponent_element,
        content=exponent,
    )

    superscript.append(
        exponent_element
    )

    return superscript


# ================================================================
# OMML FRACTION
# ================================================================

def create_fraction(
    numerator,
    denominator,
):
    """
    Create native OMML fraction.

    Numerator and denominator may contain:
        str
        OMML element
        list / tuple of OMML elements
    """

    fraction = OxmlElement(
        "m:f"
    )

    fraction_properties = OxmlElement(
        "m:fPr"
    )

    fraction.append(
        fraction_properties
    )


    # NUMERATOR

    numerator_element = OxmlElement(
        "m:num"
    )

    append_math_content(
        parent=numerator_element,
        content=numerator,
    )

    fraction.append(
        numerator_element
    )


    # DENOMINATOR

    denominator_element = OxmlElement(
        "m:den"
    )

    append_math_content(
        parent=denominator_element,
        content=denominator,
    )

    fraction.append(
        denominator_element
    )

    return fraction


# ================================================================
# OMML RADICAL
# ================================================================

def create_radical(
    expression,
):
    """
    Create native OMML square root.
    """

    radical = OxmlElement(
        "m:rad"
    )

    radical_properties = OxmlElement(
        "m:radPr"
    )

    hide_degree = OxmlElement(
        "m:degHide"
    )

    hide_degree.set(
        qn("m:val"),
        "1",
    )

    radical_properties.append(
        hide_degree
    )

    radical.append(
        radical_properties
    )


    # EMPTY DEGREE

    degree_element = OxmlElement(
        "m:deg"
    )

    radical.append(
        degree_element
    )


    # RADICAL EXPRESSION

    expression_element = OxmlElement(
        "m:e"
    )

    append_math_content(
        parent=expression_element,
        content=expression,
    )

    radical.append(
        expression_element
    )

    return radical


# ================================================================
# EQUATION PARAGRAPH
# ================================================================

def add_equation(
    document,
    equation_elements,
    equation_number=None,
):
    """
    Add native Microsoft Word equation.

    Numbered equation layout:

        left cell | centered equation | equation number

    The table borders are hidden.
    """

    if not isinstance(
        equation_elements,
        (list, tuple),
    ):

        equation_elements = [
            equation_elements
        ]


    # ============================================================
    # EQUATION WITHOUT NUMBER
    # ============================================================

    if equation_number is None:

        paragraph = document.add_paragraph()

        format_paragraph(
            paragraph=paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
        )

        math_paragraph = OxmlElement(
            "m:oMathPara"
        )

        math_object = OxmlElement(
            "m:oMath"
        )

        append_math_content(
            parent=math_object,
            content=equation_elements,
        )

        math_paragraph.append(
            math_object
        )

        paragraph._element.append(
            math_paragraph
        )

        return paragraph


    # ============================================================
    # NUMBERED EQUATION
    # ============================================================

    table = document.add_table(
        rows=1,
        cols=3,
    )

    table.autofit = False


    # ============================================================
    # HIDE TABLE BORDERS
    # ============================================================

    table_properties = table._tbl.tblPr

    table_borders = OxmlElement(
        "w:tblBorders"
    )

    for border_name in (
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ):

        border = OxmlElement(
            f"w:{border_name}"
        )

        border.set(
            qn("w:val"),
            "nil",
        )

        table_borders.append(
            border
        )

    table_properties.append(
        table_borders
    )


    # ============================================================
    # LEFT CELL
    # ============================================================

    left_cell = table.cell(
        0,
        0,
    )

    left_paragraph = (
        left_cell.paragraphs[0]
    )

    format_paragraph(
        paragraph=left_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
    )


    # ============================================================
    # EQUATION CELL
    # ============================================================

    equation_cell = table.cell(
        0,
        1,
    )

    equation_paragraph = (
        equation_cell.paragraphs[0]
    )

    format_paragraph(
        paragraph=equation_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
    )

    math_paragraph = OxmlElement(
        "m:oMathPara"
    )

    math_object = OxmlElement(
        "m:oMath"
    )

    append_math_content(
        parent=math_object,
        content=equation_elements,
    )

    math_paragraph.append(
        math_object
    )

    equation_paragraph._element.append(
        math_paragraph
    )


    # ============================================================
    # NUMBER CELL
    # ============================================================

    number_cell = table.cell(
        0,
        2,
    )

    number_paragraph = (
        number_cell.paragraphs[0]
    )

    format_paragraph(
        paragraph=number_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False,
    )

    number_run = number_paragraph.add_run(
        f"({equation_number})"
    )

    set_run_font(
        number_run
    )

    return table


# ================================================================
# DOCUMENT STYLES
# ================================================================

def configure_document_styles(
    document,
):
    """
    Configure default Word document style.
    """

    normal_style = document.styles[
        "Normal"
    ]

    normal_style.font.name = FONT_NAME

    normal_style.font.size = Pt(
        FONT_SIZE_PT
    )

    normal_style_properties = (
        normal_style._element.get_or_add_rPr()
    )

    normal_style_fonts = (
        normal_style_properties.rFonts
    )

    if normal_style_fonts is None:

        normal_style_fonts = OxmlElement(
            "w:rFonts"
        )

        normal_style_properties.insert(
            0,
            normal_style_fonts,
        )

    normal_style_fonts.set(
        qn("w:ascii"),
        FONT_NAME,
    )

    normal_style_fonts.set(
        qn("w:hAnsi"),
        FONT_NAME,
    )

    normal_style_fonts.set(
        qn("w:eastAsia"),
        FONT_NAME,
    )

    normal_style_fonts.set(
        qn("w:cs"),
        FONT_NAME,
    )


# ================================================================
# PAGE SETTINGS
# ================================================================

def configure_page_settings(
    document,
):
    """
    Configure A4 report page margins.

    Left   : 3.0 cm
    Right  : 1.5 cm
    Top    : 2.0 cm
    Bottom : 2.0 cm
    """

    section = document.sections[0]

    section.left_margin = Cm(
        3.0
    )

    section.right_margin = Cm(
        1.5
    )

    section.top_margin = Cm(
        2.0
    )

    section.bottom_margin = Cm(
        2.0
    )


# ================================================================
# NUMBER FORMATTING
# ================================================================

def format_number(
    value,
    digits=3,
):
    """
    Format number using decimal comma.
    """

    formatted = f"{value:.{digits}f}"

    return formatted.replace(
        ".",
        ",",
    )


def format_integer(
    value,
):
    """
    Format integer engineering value.
    """

    return f"{value:.0f}"


# ================================================================
# SECTION 01 REPORT
# ================================================================

def create_section01_word_report(
    result,
    P2,
    U1,
    poles,
    protection,
    f1,
    output_path=None,
):
    """
    Create Word report for Section 1.

    Parameters
    ----------
    result : dict
        Final Section 1 calculation result.

    P2 : float
        Rated output power, kW.

    U1 : float
        Rated line voltage, V.

    poles : int
        Number of poles 2p.

    protection : str
        Protection degree.

    f1 : float
        Supply frequency, Hz.

    output_path : str or Path, optional
        Custom DOCX output path.

    Returns
    -------
    pathlib.Path
        Created report path.
    """


    # ============================================================
    # OUTPUT PATH
    # ============================================================

    if output_path is None:

        REPORT_DIRECTORY.mkdir(
            parents=True,
            exist_ok=True,
        )

        output_path = (
            REPORT_DIRECTORY
            / REPORT_FILENAME
        )

    else:

        output_path = Path(
            output_path
        )

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )


    # ============================================================
    # DOCUMENT
    # ============================================================

    document = Document()

    configure_document_styles(
        document
    )

    configure_page_settings(
        document
    )


    # ============================================================
    # SECTION TITLE
    # ============================================================

    add_section_title(
        document=document,
        text="1 ВЫБОР ГЛАВНЫХ РАЗМЕРОВ",
    )


    # ============================================================
    # INITIAL DATA
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Главные размеры асинхронного двигателя "
            "выбираем в соответствии с методикой "
            "проектирования электрических машин. "
            "Исходными данными для расчета являются "
            f"номинальная полезная мощность P₂ = "
            f"{format_number(P2, 2)} кВт, номинальное "
            f"линейное напряжение U₁ = "
            f"{format_integer(U1)} В, число полюсов "
            f"2p = {poles}, степень защиты "
            f"{protection} и частота питающей сети "
            f"f₁ = {format_number(f1, 1)} Гц."
        ),
    )


    # ============================================================
    # P01 - SHAFT HEIGHT
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Высоту оси вращения асинхронного двигателя "
            "выбираем в зависимости от номинальной мощности, "
            "числа полюсов и степени защиты двигателя по "
            "рис. 9.18."
        ),
    )

    add_text_paragraph(
        document=document,
        text=(
            f"Для двигателя мощностью P₂ = "
            f"{format_number(P2, 2)} кВт, с числом полюсов "
            f"2p = {poles} и степенью защиты "
            f"{protection} принимаем высоту оси вращения:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "h="
            ),
            create_math_text(
                f"{format_integer(result['h'])} мм"
            ),
        ],
    )


    # ============================================================
    # P02 - STATOR OUTER DIAMETER
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "По принятой высоте оси вращения определяем "
            "наружный диаметр сердечника статора. "
            f"Для h = {format_integer(result['h'])} мм "
            "рекомендуемый диапазон наружного диаметра "
            "сердечника статора составляет "
            f"Dₐ = {format_number(result['Da_min'], 3)}..."
            f"{format_number(result['Da_max'], 3)} м. "
            "Принимаем:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "D",
                "a",
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['Da'], 3)} м"
            ),
        ],
    )


    # ============================================================
    # P03 - INNER DIAMETER COEFFICIENT
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Коэффициент отношения внутреннего диаметра "
            "сердечника статора к его наружному диаметру "
            "выбираем в зависимости от числа полюсов. "
            f"При 2p = {poles} рекомендуемый диапазон "
            f"составляет K_D = "
            f"{format_number(result['Kd_min'], 2)}..."
            f"{format_number(result['Kd_max'], 2)}. "
            "Принимаем:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "K",
                "D",
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['Kd'], 2)}"
            ),
        ],
    )


    # ============================================================
    # STATOR INNER DIAMETER
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Внутренний диаметр сердечника статора "
            "определяем по выражению"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "D="
            ),
            create_subscript(
                "K",
                "D",
            ),
            create_math_text(
                "·"
            ),
            create_subscript(
                "D",
                "a",
            ),
        ],
        equation_number="1.1",
    )

    add_text_paragraph(
        document=document,
        text=(
            "Подставляя численные значения, получаем"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "D="
            ),
            create_math_text(
                f"{format_number(result['Kd'], 2)}"
            ),
            create_math_text(
                "·"
            ),
            create_math_text(
                f"{format_number(result['Da'], 3)}"
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['D'], 3)} м"
            ),
        ],
    )


    # ============================================================
    # POLE PITCH
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Полюсное деление статора определяем по формуле"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "τ="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        "πD"
                    ),
                ],
                denominator=[
                    create_math_text(
                        "2p"
                    ),
                ],
            ),
        ],
        equation_number="1.2",
    )

    add_text_paragraph(
        document=document,
        text=(
            "Подставляя численные значения, получаем"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "τ="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        "π·"
                    ),
                    create_math_text(
                        f"{format_number(result['D'], 3)}"
                    ),
                ],
                denominator=[
                    create_math_text(
                        f"{poles}"
                    ),
                ],
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['tau'], 3)} м"
            ),
        ],
    )


    # ============================================================
    # P04 - KE
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Коэффициент отношения ЭДС обмотки статора "
            "к номинальному напряжению выбираем по "
            "соответствующей расчетной зависимости. "
            "Принимаем:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "K",
                "E",
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['Ke'], 3)}"
            ),
        ],
    )


    # ============================================================
    # P05 - EFFICIENCY
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Предварительное значение коэффициента полезного "
            "действия двигателя выбираем в зависимости от "
            "номинальной мощности, числа полюсов и степени "
            "защиты:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "η="
            ),
            create_math_text(
                f"{format_number(result['eta'], 3)}"
            ),
        ],
    )


    # ============================================================
    # P06 - POWER FACTOR
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Предварительное значение коэффициента мощности "
            "двигателя принимаем:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "cosφ="
            ),
            create_math_text(
                f"{format_number(result['cos_phi'], 3)}"
            ),
        ],
    )


    # ============================================================
    # CALCULATED POWER
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Расчетную мощность электрической машины "
            "определяем по выражению"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "P′="
            ),
            create_fraction(
                numerator=[
                    create_subscript(
                        "P",
                        "2",
                    ),
                    create_subscript(
                        "K",
                        "E",
                    ),
                ],
                denominator=[
                    create_math_text(
                        "η·cosφ"
                    ),
                ],
            ),
        ],
        equation_number="1.3",
    )

    add_text_paragraph(
        document=document,
        text=(
            "Подставляя численные значения, получаем"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "P′="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        f"{format_number(P2, 2)}"
                    ),
                    create_math_text(
                        "·1000·"
                    ),
                    create_math_text(
                        f"{format_number(result['Ke'], 3)}"
                    ),
                ],
                denominator=[
                    create_math_text(
                        f"{format_number(result['eta'], 3)}"
                    ),
                    create_math_text(
                        "·"
                    ),
                    create_math_text(
                        f"{format_number(result['cos_phi'], 3)}"
                    ),
                ],
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_integer(result['P_prime'])} ВА"
            ),
        ],
    )


    # ============================================================
    # P07 - AIR-GAP FLUX DENSITY
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Предварительное значение магнитной индукции "
            "в воздушном зазоре выбираем по рекомендуемому "
            "диапазону. Для рассматриваемого двигателя "
            f"Bδ = {format_number(result['B_delta_min'], 3)}..."
            f"{format_number(result['B_delta_max'], 3)} Тл. "
            "Принимаем:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "B",
                "δ",
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['B_delta'], 3)} Тл"
            ),
        ],
    )


    # ============================================================
    # P08 - LINEAR CURRENT LOADING
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Предварительное значение линейной токовой "
            "нагрузки выбираем по рекомендуемому диапазону "
            "для принятого наружного диаметра сердечника "
            "статора. Рекомендуемый диапазон составляет "
            f"A = {format_integer(result['A_min'])}..."
            f"{format_integer(result['A_max'])} А/м. "
            "Принимаем:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "A="
            ),
            create_math_text(
                f"{format_integer(result['A'])} А/м"
            ),
        ],
    )


    # ============================================================
    # ALPHA DELTA
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Коэффициент формы кривой распределения магнитной "
            "индукции в воздушном зазоре предварительно "
            "принимаем равным"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "α",
                "δ",
            ),
            create_math_text(
                "="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        "2"
                    ),
                ],
                denominator=[
                    create_math_text(
                        "π"
                    ),
                ],
            ),
            create_math_text(
                "≈"
            ),
            create_math_text(
                f"{format_number(result['alpha_delta'], 3)}"
            ),
        ],
    )


    # ============================================================
    # KB
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Коэффициент формы магнитного поля определяем "
            "по выражению"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "k",
                "B",
            ),
            create_math_text(
                "="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        "π"
                    ),
                ],
                denominator=[
                    create_math_text(
                        "2"
                    ),
                    create_radical(
                        create_math_text(
                            "2"
                        )
                    ),
                ],
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['k_B'], 3)}"
            ),
        ],
    )


    # ============================================================
    # WINDING FACTOR
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Предварительное значение обмоточного коэффициента "
            "обмотки статора принимаем:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "k",
                "об1",
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['kw1'], 3)}"
            ),
        ],
    )


    # ============================================================
    # SYNCHRONOUS SPEED
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Синхронная частота вращения магнитного поля "
            "статора составляет:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "n",
                "1",
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_integer(result['n1_rpm'])} об/мин"
            ),
        ],
    )


    # ============================================================
    # SYNCHRONOUS ANGULAR FREQUENCY
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Синхронную угловую частоту вращения магнитного "
            "поля определяем по формуле"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "Ω="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        "2π"
                    ),
                    create_subscript(
                        "f",
                        "1",
                    ),
                ],
                denominator=[
                    create_math_text(
                        "p"
                    ),
                ],
            ),
        ],
        equation_number="1.4",
    )

    pole_pairs = (
        poles
        / 2
    )

    add_text_paragraph(
        document=document,
        text=(
            "Подставляя численные значения, получаем"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "Ω="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        "2π·"
                    ),
                    create_math_text(
                        f"{format_number(f1, 1)}"
                    ),
                ],
                denominator=[
                    create_math_text(
                        f"{format_number(pole_pairs, 0)}"
                    ),
                ],
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['Omega'], 3)} рад/с"
            ),
        ],
    )


    # ============================================================
    # CALCULATED CORE LENGTH
    # KOPYLOV FORMULA 9.6
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Расчетную длину магнитопровода статора определяем "
            "по формуле"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "l",
                "δ",
            ),
            create_math_text(
                "="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        "P′"
                    ),
                ],
                denominator=[
                    create_superscript(
                        "D",
                        "2",
                    ),
                    create_math_text(
                        "Ω"
                    ),
                    create_subscript(
                        "k",
                        "B",
                    ),
                    create_subscript(
                        "k",
                        "об1",
                    ),
                    create_math_text(
                        "A"
                    ),
                    create_subscript(
                        "B",
                        "δ",
                    ),
                ],
            ),
        ],
        equation_number="1.5",
    )

    add_text_paragraph(
        document=document,
        text=(
            "Подставляя полученные значения параметров, "
            "определяем расчетную длину магнитопровода:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_subscript(
                "l",
                "δ",
            ),
            create_math_text(
                "="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        f"{format_integer(result['P_prime'])}"
                    ),
                ],
                denominator=[
    create_superscript(
        f"{format_number(result['D'], 4)}",
        "2",
    ),
    create_math_text(
        "·"
    ),
    create_math_text(
        f"{format_number(result['Omega'], 3)}"
    ),
    create_math_text(
        "·"
    ),
    create_math_text(
        f"{format_number(result['k_B'], 4)}"
    ),
    create_math_text(
        "·"
    ),
    create_math_text(
        f"{format_number(result['kw1'], 3)}"
    ),
    create_math_text(
        "·"
    ),
    create_math_text(
        f"{format_integer(result['A'])}"
    ),
    create_math_text(
        "·"
    ),
    create_math_text(
        f"{format_number(result['B_delta'], 3)}"
    ),
],
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['l_delta'], 4)} м"
            ),
        ],
    )


    # ============================================================
    # P09 - LAMBDA CHECK
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Правильность выбора главных размеров двигателя "
            "проверяем по отношению расчетной длины "
            "магнитопровода к полюсному делению:"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "λ="
            ),
            create_fraction(
                numerator=[
                    create_subscript(
                        "l",
                        "δ",
                    ),
                ],
                denominator=[
                    create_math_text(
                        "τ"
                    ),
                ],
            ),
        ],
        equation_number="1.6",
    )

    add_text_paragraph(
        document=document,
        text=(
            "Для проектируемого двигателя получаем"
        ),
    )

    add_equation(
        document=document,
        equation_elements=[
            create_math_text(
                "λ="
            ),
            create_fraction(
                numerator=[
                    create_math_text(
                        f"{format_number(result['l_delta'], 4)}"
                    ),
                ],
                denominator=[
                    create_math_text(
                        f"{format_number(result['tau'], 4)}"
                    ),
                ],
            ),
            create_math_text(
                "="
            ),
            create_math_text(
                f"{format_number(result['lambda'], 3)}"
            ),
        ],
    )


    # ============================================================
    # LAMBDA CONCLUSION
    # ============================================================

    add_text_paragraph(
        document=document,
        text=(
            "Согласно рекомендациям методики проектирования "
            "допустимый диапазон отношения λ для выбранного "
            "исполнения двигателя составляет "
            f"{format_number(result['lambda_min'], 3)}..."
            f"{format_number(result['lambda_max'], 3)}."
        ),
    )

    if result["status"] == "PERMISSIBLE":

        add_text_paragraph(
            document=document,
            text=(
                f"Полученное значение λ = "
                f"{format_number(result['lambda'], 3)} "
                "находится в допустимом диапазоне. "
                "Следовательно, выбранные главные размеры "
                "асинхронного двигателя удовлетворяют "
                "рекомендациям методики проектирования."
            ),
        )

    else:

        add_text_paragraph(
            document=document,
            text=(
                f"Полученное значение λ = "
                f"{format_number(result['lambda'], 3)} "
                "не находится в допустимом диапазоне. "
                "Выбранные главные размеры требуют "
                "дополнительной корректировки."
            ),
        )


    # ============================================================
    # SAVE DOCUMENT
    # ============================================================

    document.save(
        output_path
    )

    return output_path


# ================================================================
# MODULE INFORMATION
# ================================================================

if __name__ == "__main__":

    print(
        "DEE Section 01 Word Report Generator"
    )

    print(
        "Native OMML subscripts and superscripts enabled."
    )

    print(
        "This module must be called from main.py "
        "with the final Section 1 calculation result."
    )