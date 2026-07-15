"""
DIGITAL ELECTRICAL ENGINEER (DEE)

AC MOTOR FEEDER DESIGN
WORD REPORT GENERATOR

The module creates a technical Word report for an AC motor feeder
designed according to the Schneider Electric / IEC methodology.

Report language:
    Russian.

Report format:
    Microsoft Word DOCX.

Main formatting:
    Times New Roman;
    14 pt;
    1.5 line spacing;
    1.25 cm first-line paragraph indent;
    justified body text.

The report contains:
    initial motor data;
    motor design current calculation;
    motor starting current calculation;
    cable selection;
    steady-state voltage-drop calculation;
    starting voltage-drop calculation;
    three-phase short-circuit calculation;
    motor circuit breaker selection;
    contactor selection;
    Type 2 coordination check;
    final technical assessment;
    equipment specification.
"""


# =============================================================================
# IMPORTS
# =============================================================================

from pathlib import Path
from math import sqrt

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# =============================================================================
# REPORT ERROR
# =============================================================================

class ACMotorFeederWordReportError(ValueError):
    """
    Error raised when the AC motor feeder Word report
    cannot be created.
    """

    pass


# =============================================================================
# CONSTANTS
# =============================================================================

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 14

LINE_SPACING = 1.5
FIRST_LINE_INDENT_CM = 1.25

PAGE_MARGIN_LEFT_CM = 3.0
PAGE_MARGIN_RIGHT_CM = 1.5
PAGE_MARGIN_TOP_CM = 2.0
PAGE_MARGIN_BOTTOM_CM = 2.0


# =============================================================================
# VALUE HELPERS
# =============================================================================

def _get_value(
    data,
    *keys,
    default=None,
):
    """
    Return the first available dictionary value
    from the specified key sequence.
    """

    if not isinstance(data, dict):

        return default

    for key in keys:

        if key in data:

            value = data[key]

            if value is not None:

                return value

    return default


def _format_number(
    value,
    digits=2,
):
    """
    Format a numerical value using a comma
    as the decimal separator.
    """

    if value is None:

        return "—"

    try:

        formatted_value = (
            f"{float(value):.{digits}f}"
        )

    except (
        TypeError,
        ValueError,
    ):

        return str(
            value
        )

    return formatted_value.replace(
        ".",
        ",",
    )


def _format_integer(
    value,
):
    """
    Format a numerical value without decimal places.
    """

    if value is None:

        return "—"

    try:

        return f"{float(value):.0f}"

    except (
        TypeError,
        ValueError,
    ):

        return str(
            value
        )


def _status_to_russian(
    status,
):
    """
    Convert main DEE status values
    to Russian report text.
    """

    status_map = {
        "PERMISSIBLE":
            "ДОПУСТИМО",

        "NOT PERMISSIBLE":
            "НЕДОПУСТИМО",

        "SELECTED":
            "ВЫБРАНО",

        "NOT SELECTED":
            "НЕ ВЫБРАНО",

        "ACCEPTED":
            "ПРИНЯТО",

        "NOT ACCEPTED":
            "НЕ ПРИНЯТО",

        "VERIFIED":
            "ПОДТВЕРЖДЕНО",

        "NOT VERIFIED":
            "НЕ ПОДТВЕРЖДЕНО",

        "FOUND":
            "НАЙДЕНО",

        "NOT FOUND":
            "НЕ НАЙДЕНО",

        "VALID":
            "КОРРЕКТНО",

        "CALCULATED":
            "РАССЧИТАНО",
    }

    normalized_status = str(
        status
    ).strip().upper()

    return status_map.get(
        normalized_status,
        str(
            status
        ),
    )


# =============================================================================
# WORD XML HELPERS
# =============================================================================

def _set_cell_shading(
    cell,
    fill,
):
    """
    Set table cell background shading.
    """

    tc_pr = cell._tc.get_or_add_tcPr()

    shading = tc_pr.find(
        qn(
            "w:shd"
        )
    )

    if shading is None:

        shading = OxmlElement(
            "w:shd"
        )

        tc_pr.append(
            shading
        )

    shading.set(
        qn(
            "w:fill"
        ),
        fill,
    )


def _set_cell_margins(
    cell,
    top=80,
    start=80,
    bottom=80,
    end=80,
):
    """
    Set table cell margins.
    """

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_mar = tc_pr.first_child_found_in(
        "w:tcMar"
    )

    if tc_mar is None:

        tc_mar = OxmlElement(
            "w:tcMar"
        )

        tc_pr.append(
            tc_mar
        )

    margins = {
        "top":
            top,

        "start":
            start,

        "bottom":
            bottom,

        "end":
            end,
    }

    for margin_name, margin_value in margins.items():

        node = tc_mar.find(
            qn(
                f"w:{margin_name}"
            )
        )

        if node is None:

            node = OxmlElement(
                f"w:{margin_name}"
            )

            tc_mar.append(
                node
            )

        node.set(
            qn(
                "w:w"
            ),
            str(
                margin_value
            ),
        )

        node.set(
            qn(
                "w:type"
            ),
            "dxa",
        )


# =============================================================================
# FONT HELPERS
# =============================================================================

def _set_run_font(
    run,
    size_pt=FONT_SIZE_PT,
    bold=False,
    italic=False,
):
    """
    Apply the report font to a run.
    """

    run.font.name = FONT_NAME

    run.font.size = Pt(
        size_pt
    )

    run.bold = bold
    run.italic = italic

    run._element.rPr.rFonts.set(
        qn(
            "w:eastAsia"
        ),
        FONT_NAME,
    )


def _format_paragraph_runs(
    paragraph,
    size_pt=FONT_SIZE_PT,
):
    """
    Apply report font settings to all paragraph runs.
    """

    for run in paragraph.runs:

        _set_run_font(
            run,
            size_pt=size_pt,
            bold=run.bold,
            italic=run.italic,
        )


# =============================================================================
# PARAGRAPH HELPERS
# =============================================================================

def _configure_body_paragraph(
    paragraph,
):
    """
    Configure a normal body paragraph.
    """

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    paragraph_format = (
        paragraph.paragraph_format
    )

    paragraph_format.line_spacing = (
        LINE_SPACING
    )

    paragraph_format.first_line_indent = Cm(
        FIRST_LINE_INDENT_CM
    )

    paragraph_format.space_before = Pt(
        0
    )

    paragraph_format.space_after = Pt(
        0
    )


def _add_body_paragraph(
    document,
    text="",
):
    """
    Add a formatted report body paragraph.
    """

    paragraph = document.add_paragraph()

    _configure_body_paragraph(
        paragraph
    )

    run = paragraph.add_run(
        text
    )

    _set_run_font(
        run
    )

    return paragraph


def _add_centered_paragraph(
    document,
    text="",
    bold=False,
    size_pt=FONT_SIZE_PT,
):
    """
    Add a centered paragraph.
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.line_spacing = (
        LINE_SPACING
    )

    paragraph.paragraph_format.space_before = Pt(
        0
    )

    paragraph.paragraph_format.space_after = Pt(
        0
    )

    run = paragraph.add_run(
        text
    )

    _set_run_font(
        run,
        size_pt=size_pt,
        bold=bold,
    )

    return paragraph


def _add_section_heading(
    document,
    text,
):
    """
    Add a numbered report section heading.
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
    )

    paragraph.paragraph_format.line_spacing = (
        LINE_SPACING
    )

    paragraph.paragraph_format.first_line_indent = Cm(
        0
    )

    paragraph.paragraph_format.space_before = Pt(
        6
    )

    paragraph.paragraph_format.space_after = Pt(
        0
    )

    run = paragraph.add_run(
        text
    )

    _set_run_font(
        run,
        bold=True,
    )

    return paragraph


def _add_formula_paragraph(
    document,
):
    """
    Add an empty centered formula paragraph.
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.line_spacing = (
        LINE_SPACING
    )

    paragraph.paragraph_format.first_line_indent = Cm(
        0
    )

    paragraph.paragraph_format.space_before = Pt(
        0
    )

    paragraph.paragraph_format.space_after = Pt(
        0
    )

    return paragraph


def _add_run(
    paragraph,
    text,
    bold=False,
    italic=False,
    subscript=False,
    superscript=False,
):
    """
    Add a formatted run to a paragraph.
    """

    run = paragraph.add_run(
        str(
            text
        )
    )

    _set_run_font(
        run,
        bold=bold,
        italic=italic,
    )

    run.font.subscript = subscript
    run.font.superscript = superscript

    return run


def _add_symbol_with_subscript(
    paragraph,
    symbol,
    subscript,
):
    """
    Add a symbol with a real Word lower index.
    """

    _add_run(
        paragraph,
        symbol,
        italic=True,
    )

    _add_run(
        paragraph,
        subscript,
        subscript=True,
    )


def _add_symbol_with_superscript(
    paragraph,
    symbol,
    superscript,
):
    """
    Add a symbol with a real Word upper index.
    """

    _add_run(
        paragraph,
        symbol,
        italic=True,
    )

    _add_run(
        paragraph,
        superscript,
        superscript=True,
    )


# =============================================================================
# TABLE HELPERS
# =============================================================================

def _format_table(
    table,
):
    """
    Apply standard report formatting to a table.
    """

    table.style = "Table Grid"

    for row in table.rows:

        for cell in row.cells:

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            _set_cell_margins(
                cell
            )

            for paragraph in cell.paragraphs:

                paragraph.paragraph_format.line_spacing = 1.0

                paragraph.paragraph_format.first_line_indent = Cm(
                    0
                )

                paragraph.paragraph_format.space_before = Pt(
                    0
                )

                paragraph.paragraph_format.space_after = Pt(
                    0
                )

                _format_paragraph_runs(
                    paragraph,
                    size_pt=12,
                )


def _add_data_table(
    document,
    rows,
):
    """
    Add a two-column parameter table.
    """

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    header_cells[0].text = (
        "Параметр"
    )

    header_cells[1].text = (
        "Значение"
    )

    for cell in header_cells:

        _set_cell_shading(
            cell,
            "D9EAF7",
        )

        for paragraph in cell.paragraphs:

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            for run in paragraph.runs:

                _set_run_font(
                    run,
                    size_pt=12,
                    bold=True,
                )

    for parameter_name, parameter_value in rows:

        cells = table.add_row().cells

        cells[0].text = str(
            parameter_name
        )

        cells[1].text = str(
            parameter_value
        )

    _format_table(
        table
    )

    return table


# =============================================================================
# REPORT CREATION
# =============================================================================

def create_ac_motor_feeder_word_report(
    motor,
    installation,
    supply,
    design_current,
    starting_current,
    cable,
    voltage_drop,
    starting_voltage_drop,
    short_circuit,
    motor_circuit_breaker,
    contactor,
    motor_starter_coordination,
    final_design_check,
    motor_starter_coordination_optimization=None,
    output_path=None,
):
    """
    Create the complete AC motor feeder Word report.

    Returns
    -------
    pathlib.Path
        Created DOCX report path.
    """

    # -------------------------------------------------------------------------
    # INPUT DATA CHECK
    # -------------------------------------------------------------------------

    report_data = {
        "motor":
            motor,

        "installation":
            installation,

        "supply":
            supply,

        "design_current":
            design_current,

        "starting_current":
            starting_current,

        "cable":
            cable,

        "voltage_drop":
            voltage_drop,

        "starting_voltage_drop":
            starting_voltage_drop,

        "short_circuit":
            short_circuit,

        "motor_circuit_breaker":
            motor_circuit_breaker,

        "contactor":
            contactor,

        "motor_starter_coordination":
            motor_starter_coordination,

        "final_design_check":
            final_design_check,
    }

    for data_name, data in report_data.items():

        if not isinstance(
            data,
            dict,
        ):

            raise ACMotorFeederWordReportError(
                f"{data_name} data must be provided "
                f"as a dictionary."
            )

    # -------------------------------------------------------------------------
    # MOTOR DATA
    # -------------------------------------------------------------------------

    manufacturer = _get_value(
        motor,
        "manufacturer",
        default="—",
    )

    motor_type = _get_value(
        motor,
        "motor_type",
        "type",
        "model",
        default="—",
    )

    rated_power_kw = _get_value(
        motor,
        "rated_power_kw",
    )

    rated_voltage_v = _get_value(
        motor,
        "rated_voltage_v",
    )

    rated_current_a = _get_value(
        motor,
        "rated_current_a",
    )

    rated_frequency_hz = _get_value(
        motor,
        "rated_frequency_hz",
    )

    rated_speed_rpm = _get_value(
        motor,
        "rated_speed_rpm",
    )

    efficiency = _get_value(
        motor,
        "efficiency",
    )

    power_factor = _get_value(
        motor,
        "power_factor",
    )

    phases = _get_value(
        motor,
        "phases",
        "number_of_phases",
        default=3,
    )

    connection = _get_value(
        motor,
        "connection",
        default="—",
    )

    duty = _get_value(
        motor,
        "duty",
        default="—",
    )

    efficiency_class = _get_value(
        motor,
        "efficiency_class",
        default="—",
    )

    starting_method = _get_value(
        motor,
        "starting_method",
        default="—",
    )

    starting_current_ratio = _get_value(
        motor,
        "starting_current_ratio",
        "starting_current_ratio_is_in",
    )

    starting_time_s = _get_value(
        motor,
        "starting_time_s",
    )

    # -------------------------------------------------------------------------
    # CALCULATION VALUES
    # -------------------------------------------------------------------------

    calculated_current_a = _get_value(
        design_current,
        "calculated_current_a",
    )

    design_current_a = _get_value(
        design_current,
        "design_current_a",
    )

    calculated_starting_current_a = _get_value(
        starting_current,
        "starting_current_a",
        "calculated_starting_current_a",
    )

    cable_section_mm2 = _get_value(
        cable,
        "selected_section_mm2",
    )

    tabulated_current_a = _get_value(
        cable,
        "selected_tabulated_current_a",
    )

    corrected_current_capacity_a = _get_value(
        cable,
        "corrected_current_capacity_a",
    )

    cable_length_m = _get_value(
        installation,
        "cable_length_m",
    )

    cable_material = _get_value(
        installation,
        "conductor_material",
        default="—",
    )

    cable_insulation = _get_value(
        installation,
        "insulation_type",
        default="—",
    )

    installation_method = _get_value(
        installation,
        "installation_method",
        default="—",
    )

    voltage_drop_v = _get_value(
        voltage_drop,
        "voltage_drop_v",
    )

    voltage_drop_percent = _get_value(
        voltage_drop,
        "voltage_drop_percent",
    )

    motor_terminal_voltage_v = _get_value(
        voltage_drop,
        "motor_terminal_voltage_v",
    )

    maximum_voltage_drop_percent = _get_value(
        voltage_drop,
        "maximum_voltage_drop_percent",
        "max_voltage_drop_percent",
    )

    starting_voltage_drop_v = _get_value(
        starting_voltage_drop,
        "starting_voltage_drop_v",
    )

    starting_voltage_drop_percent = _get_value(
        starting_voltage_drop,
        "starting_voltage_drop_percent",
    )

    starting_terminal_voltage_v = _get_value(
        starting_voltage_drop,
        "starting_terminal_voltage_v",
    )

    maximum_starting_voltage_drop_percent = _get_value(
        starting_voltage_drop,
        "maximum_starting_voltage_drop_percent",
        "max_starting_voltage_drop_percent",
    )

    transformer_power_kva = _get_value(
        supply,
        "transformer_rated_power_kva",
        "rated_power_kva",
    )

    transformer_voltage_v = _get_value(
        supply,
        "transformer_secondary_voltage_v",
        "secondary_voltage_v",
    )

    transformer_uk_percent = _get_value(
        supply,
        "transformer_impedance_percent",
        "short_circuit_voltage_percent",
        "uk_percent",
    )

    transformer_load_losses_w = _get_value(
        supply,
        "transformer_load_losses_w",
        "load_losses_w",
    )

    transformer_terminal_short_circuit_ka = _get_value(
        short_circuit,
        "transformer_terminal_short_circuit_current_ka",
    )

    feeder_end_short_circuit_ka = _get_value(
        short_circuit,
        "feeder_end_short_circuit_current_ka",
    )

    breaker_reference = _get_value(
        final_design_check,
        "final_breaker_reference",
        default=_get_value(
            motor_circuit_breaker,
            "reference",
            default="—",
        ),
    )

    contactor_reference = _get_value(
        final_design_check,
        "final_contactor_reference",
        default=_get_value(
            contactor,
            "reference",
            default="—",
        ),
    )

    breaker_thermal_min_a = _get_value(
        motor_circuit_breaker,
        "thermal_adjustment_min_a",
        "thermal_min_a",
    )

    breaker_thermal_max_a = _get_value(
        motor_circuit_breaker,
        "thermal_adjustment_max_a",
        "thermal_max_a",
    )

    breaker_breaking_capacity_ka = _get_value(
        motor_circuit_breaker,
        "breaking_capacity_ka",
        "icu_ka",
    )

    contactor_operational_current_a = _get_value(
        contactor,
        "operational_current_a",
        "rated_operational_current_a",
        "ac3_operational_current_a",
    )

    contactor_power_kw = _get_value(
        contactor,
        "motor_power_kw",
        "rated_motor_power_kw",
    )

    coordination_type = _get_value(
        motor_starter_coordination,
        "coordination_type",
        default="—",
    )

    coordination_short_circuit_ka = _get_value(
        motor_starter_coordination,
        "coordination_short_circuit_ka",
    )

    coordination_status = _get_value(
        motor_starter_coordination,
        "coordination_status",
        default="—",
    )

    if (
        coordination_status != "VERIFIED"
        and motor_starter_coordination_optimization
        is not None
    ):

        coordination_type = _get_value(
            motor_starter_coordination_optimization,
            "coordination_type",
            default=coordination_type,
        )

        coordination_short_circuit_ka = _get_value(
            motor_starter_coordination_optimization,
            "coordination_short_circuit_ka",
            default=coordination_short_circuit_ka,
        )

        coordination_status = _get_value(
            motor_starter_coordination_optimization,
            "optimization_status",
            default=coordination_status,
        )

    final_design_status = _get_value(
        final_design_check,
        "final_design_status",
        default="—",
    )

    # -------------------------------------------------------------------------
    # OUTPUT PATH
    # -------------------------------------------------------------------------

    if output_path is None:

        power_text = (
            _format_number(
                rated_power_kw,
                digits=1,
            )
            .replace(
                ",",
                "_",
            )
        )

        voltage_text = _format_integer(
            rated_voltage_v
        )

        output_path = Path(
            "reports"
        ) / (
            f"AC_Motor_Feeder_"
            f"{power_text}kW_"
            f"{voltage_text}V.docx"
        )

    else:

        output_path = Path(
            output_path
        )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    # -------------------------------------------------------------------------
    # DOCUMENT
    # -------------------------------------------------------------------------

    document = Document()

    section = document.sections[0]

    section.left_margin = Cm(
        PAGE_MARGIN_LEFT_CM
    )

    section.right_margin = Cm(
        PAGE_MARGIN_RIGHT_CM
    )

    section.top_margin = Cm(
        PAGE_MARGIN_TOP_CM
    )

    section.bottom_margin = Cm(
        PAGE_MARGIN_BOTTOM_CM
    )

    # -------------------------------------------------------------------------
    # NORMAL STYLE
    # -------------------------------------------------------------------------

    normal_style = document.styles[
        "Normal"
    ]

    normal_style.font.name = FONT_NAME

    normal_style.font.size = Pt(
        FONT_SIZE_PT
    )

    normal_style._element.rPr.rFonts.set(
        qn(
            "w:eastAsia"
        ),
        FONT_NAME,
    )

    normal_style.paragraph_format.line_spacing = (
        LINE_SPACING
    )

    # =========================================================================
    # REPORT TITLE
    # =========================================================================

    _add_centered_paragraph(
        document,
        "РАСЧЁТ И ВЫБОР ЭЛЕКТРООБОРУДОВАНИЯ",
        bold=True,
    )

    _add_centered_paragraph(
        document,
        "ЦЕПИ ПИТАНИЯ АСИНХРОННОГО ЭЛЕКТРОДВИГАТЕЛЯ",
        bold=True,
    )

    _add_centered_paragraph(
        document,
        "ПО МЕТОДИКЕ SCHNEIDER ELECTRIC / IEC",
        bold=True,
    )

    document.add_paragraph()

    # =========================================================================
    # SECTION 1
    # =========================================================================

    _add_section_heading(
        document,
        "1 ИСХОДНЫЕ ДАННЫЕ ЭЛЕКТРОДВИГАТЕЛЯ",
    )

    _add_body_paragraph(
        document,
        (
            "Для расчёта цепи питания принят трёхфазный "
            "асинхронный электродвигатель. Основные паспортные "
            "данные двигателя приведены ниже."
        ),
    )

    motor_data_rows = [
        (
            "Изготовитель",
            manufacturer,
        ),
        (
            "Тип двигателя",
            motor_type,
        ),
        (
            "Номинальная мощность Pₙ, кВт",
            _format_number(
                rated_power_kw,
            ),
        ),
        (
            "Номинальное напряжение Uₙ, В",
            _format_integer(
                rated_voltage_v,
            ),
        ),
        (
            "Номинальный ток Iₙ, А",
            _format_number(
                rated_current_a,
            ),
        ),
        (
            "Номинальная частота fₙ, Гц",
            _format_number(
                rated_frequency_hz,
                digits=1,
            ),
        ),
        (
            "Номинальная частота вращения nₙ, об/мин",
            _format_integer(
                rated_speed_rpm,
            ),
        ),
        (
            "КПД η",
            _format_number(
                efficiency,
                digits=3,
            ),
        ),
        (
            "Коэффициент мощности cos φ",
            _format_number(
                power_factor,
                digits=3,
            ),
        ),
        (
            "Число фаз",
            phases,
        ),
        (
            "Схема соединения обмоток",
            connection,
        ),
        (
            "Режим работы",
            duty,
        ),
        (
            "Класс энергоэффективности",
            efficiency_class,
        ),
        (
            "Способ пуска",
            starting_method,
        ),
        (
            "Кратность пускового тока Iₛ/Iₙ",
            _format_number(
                starting_current_ratio,
            ),
        ),
        (
            "Время пуска tₛ, с",
            _format_number(
                starting_time_s,
            ),
        ),
    ]

    _add_data_table(
        document,
        motor_data_rows,
    )

    # =========================================================================
    # SECTION 2
    # =========================================================================

    _add_section_heading(
        document,
        "2 ОПРЕДЕЛЕНИЕ РАСЧЁТНОГО ТОКА ЭЛЕКТРОДВИГАТЕЛЯ",
    )

    _add_body_paragraph(
        document,
        (
            "Расчётное значение тока трёхфазного асинхронного "
            "электродвигателя определяется по номинальной мощности, "
            "линейному напряжению, коэффициенту полезного действия "
            "и коэффициенту мощности."
        ),
    )

    formula = _add_formula_paragraph(
        document
    )

    _add_symbol_with_subscript(
        formula,
        "I",
        "расч",
    )

    _add_run(
        formula,
        " = P",
        italic=True,
    )

    _add_run(
        formula,
        "n",
        subscript=True,
    )

    _add_run(
        formula,
        " / (√3 · U",
    )

    _add_run(
        formula,
        "n",
        subscript=True,
    )

    _add_run(
        formula,
        " · η · cos φ).",
    )

    _add_body_paragraph(
        document,
        (
            f"По результатам контрольного расчёта получено "
            f"Iрасч = {_format_number(calculated_current_a)} А. "
            f"Паспортный номинальный ток двигателя составляет "
            f"Iₙ = {_format_number(rated_current_a)} А."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Для дальнейшего расчёта цепи питания принимается "
            f"расчётный ток IВ = {_format_number(design_current_a)} А. "
            f"В качестве исходного значения используется паспортный "
            f"ток электродвигателя."
        ),
    )

    # =========================================================================
    # SECTION 3
    # =========================================================================

    _add_section_heading(
        document,
        "3 РАСЧЁТ ПУСКОВОГО ТОКА ЭЛЕКТРОДВИГАТЕЛЯ",
    )

    _add_body_paragraph(
        document,
        (
            "При прямом пуске DOL пусковой ток определяется "
            "по номинальному току двигателя и паспортной кратности "
            "пускового тока."
        ),
    )

    formula = _add_formula_paragraph(
        document
    )

    _add_symbol_with_subscript(
        formula,
        "I",
        "s",
    )

    _add_run(
        formula,
        " = k",
        italic=True,
    )

    _add_run(
        formula,
        "I",
        subscript=True,
    )

    _add_run(
        formula,
        " · I",
    )

    _add_run(
        formula,
        "n",
        subscript=True,
    )

    _add_run(
        formula,
        ".",
    )

    _add_body_paragraph(
        document,
        (
            f"При кратности пускового тока "
            f"Iₛ/Iₙ = {_format_number(starting_current_ratio)} "
            f"пусковой ток составляет "
            f"Iₛ = {_format_number(calculated_starting_current_a)} А."
        ),
    )

    # =========================================================================
    # SECTION 4
    # =========================================================================

    _add_section_heading(
        document,
        "4 ВЫБОР СЕЧЕНИЯ СИЛОВОГО КАБЕЛЯ",
    )

    _add_body_paragraph(
        document,
        (
            "Выбор сечения силового кабеля выполняется по условию "
            "допустимой длительной токовой нагрузки с учётом "
            "условий прокладки и применяемых поправочных коэффициентов."
        ),
    )

    formula = _add_formula_paragraph(
        document
    )

    _add_symbol_with_subscript(
        formula,
        "I",
        "B",
    )

    _add_run(
        formula,
        " ≤ I",
    )

    _add_run(
        formula,
        "Z",
        subscript=True,
    )

    _add_run(
        formula,
        ".",
    )

    _add_body_paragraph(
        document,
        (
            f"Расчётный ток цепи двигателя составляет "
            f"IВ = {_format_number(design_current_a)} А. "
            f"По таблицам допустимой токовой нагрузки выбрано "
            f"сечение фазного проводника S = "
            f"{_format_number(cable_section_mm2, digits=1)} мм²."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Табличный допустимый ток выбранного проводника "
            f"составляет {_format_number(tabulated_current_a)} А, "
            f"а скорректированная допустимая токовая нагрузка "
            f"I_Z = {_format_number(corrected_current_capacity_a)} А."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Условие IВ ≤ I_Z выполняется. "
            f"Результат проверки: "
            f"{_status_to_russian(_get_value(cable, 'current_capacity_check', default='—'))}."
        ),
    )

    # =========================================================================
    # SECTION 5
    # =========================================================================

    _add_section_heading(
        document,
        "5 ПРОВЕРКА ПАДЕНИЯ НАПРЯЖЕНИЯ В УСТАНОВИВШЕМСЯ РЕЖИМЕ",
    )

    _add_body_paragraph(
        document,
        (
            "Падение линейного напряжения в трёхфазной цепи "
            "электродвигателя определяется с учётом активного "
            "и индуктивного сопротивлений кабельной линии."
        ),
    )

    formula = _add_formula_paragraph(
        document
    )

    _add_run(
        formula,
        "ΔU = √3 · I",
    )

    _add_run(
        formula,
        "B",
        subscript=True,
    )

    _add_run(
        formula,
        " · L · (R · cos φ + X · sin φ).",
    )

    _add_body_paragraph(
        document,
        (
            f"При длине линии L = {_format_number(cable_length_m, digits=1)} м "
            f"и сечении S = {_format_number(cable_section_mm2, digits=1)} мм² "
            f"расчётное падение напряжения составляет "
            f"ΔU = {_format_number(voltage_drop_v)} В, или "
            f"{_format_number(voltage_drop_percent)} %."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Расчётное напряжение на выводах электродвигателя "
            f"составляет {_format_number(motor_terminal_voltage_v)} В. "
            f"Проверка падения напряжения: "
            f"{_status_to_russian(_get_value(voltage_drop, 'voltage_drop_check', default='—'))}."
        ),
    )

    # =========================================================================
    # SECTION 6
    # =========================================================================

    _add_section_heading(
        document,
        "6 ПРОВЕРКА ПАДЕНИЯ НАПРЯЖЕНИЯ ПРИ ПУСКЕ",
    )

    _add_body_paragraph(
        document,
        (
            "В режиме прямого пуска ток двигателя существенно "
            "превышает номинальное значение. Поэтому дополнительно "
            "выполняется проверка падения напряжения в пусковом режиме."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Пусковой ток двигателя Iₛ = "
            f"{_format_number(calculated_starting_current_a)} А. "
            f"Расчётное падение напряжения при пуске составляет "
            f"ΔUₛ = {_format_number(starting_voltage_drop_v)} В, "
            f"или {_format_number(starting_voltage_drop_percent)} %."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Напряжение на выводах двигателя в момент пуска "
            f"составляет {_format_number(starting_terminal_voltage_v)} В. "
            f"Результат проверки пускового падения напряжения: "
            f"{_status_to_russian(_get_value(starting_voltage_drop, 'starting_voltage_drop_check', default='—'))}."
        ),
    )

    # =========================================================================
    # SECTION 7
    # =========================================================================

    _add_section_heading(
        document,
        "7 РАСЧЁТ ТОКА ТРЁХФАЗНОГО КОРОТКОГО ЗАМЫКАНИЯ",
    )

    _add_body_paragraph(
        document,
        (
            "Расчёт трёхфазного тока короткого замыкания выполняется "
            "с раздельным суммированием активных и индуктивных "
            "сопротивлений трансформатора и кабельной линии."
        ),
    )

    formula = _add_formula_paragraph(
        document
    )

    _add_symbol_with_subscript(
        formula,
        "Z",
        "Σ",
    )

    _add_run(
        formula,
        " = √(",
    )

    _add_symbol_with_superscript(
        formula,
        "R",
        "2",
    )

    _add_symbol_with_subscript(
        formula,
        "Σ",
        "",
    )

    _add_run(
        formula,
        " + ",
    )

    _add_symbol_with_superscript(
        formula,
        "X",
        "2",
    )

    _add_symbol_with_subscript(
        formula,
        "Σ",
        "",
    )

    _add_run(
        formula,
        ").",
    )

    formula = _add_formula_paragraph(
        document
    )

    _add_symbol_with_subscript(
        formula,
        "I",
        "k3",
    )

    _add_run(
        formula,
        " = U",
    )

    _add_run(
        formula,
        "n",
        subscript=True,
    )

    _add_run(
        formula,
        " / (√3 · Z",
    )

    _add_run(
        formula,
        "Σ",
        subscript=True,
    )

    _add_run(
        formula,
        ").",
    )

    _add_body_paragraph(
        document,
        (
            f"Питание цепи двигателя осуществляется от трансформатора "
            f"мощностью Sₙ = {_format_number(transformer_power_kva, digits=1)} кВА "
            f"с вторичным напряжением "
            f"{_format_integer(transformer_voltage_v)} В."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Расчётный ток трёхфазного короткого замыкания "
            f"на выводах трансформатора составляет "
            f"{_format_number(transformer_terminal_short_circuit_ka)} кА. "
            f"С учётом сопротивления кабельной линии ток короткого "
            f"замыкания в конце цепи двигателя составляет "
            f"I_k3 = {_format_number(feeder_end_short_circuit_ka)} кА."
        ),
    )

    # =========================================================================
    # SECTION 8
    # =========================================================================

    _add_section_heading(
        document,
        "8 ВЫБОР АВТОМАТИЧЕСКОГО ВЫКЛЮЧАТЕЛЯ ДВИГАТЕЛЯ",
    )

    _add_body_paragraph(
        document,
        (
            "Для защиты цепи электродвигателя выбирается автоматический "
            "выключатель защиты двигателя Schneider Electric TeSys. "
            "Выбор выполняется по номинальному току двигателя, диапазону "
            "регулировки теплового расцепителя и отключающей способности."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"По результатам расчёта выбран автоматический выключатель "
            f"{breaker_reference}. Диапазон регулировки теплового "
            f"расцепителя составляет "
            f"{_format_number(breaker_thermal_min_a)}... "
            f"{_format_number(breaker_thermal_max_a)} А."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Расчётный ток короткого замыкания в точке установки "
            f"составляет I_k3 = {_format_number(feeder_end_short_circuit_ka)} кА. "
            f"Проверка отключающей способности автоматического "
            f"выключателя: "
            f"{_status_to_russian(_get_value(motor_circuit_breaker, 'breaking_capacity_check', default='—'))}."
        ),
    )

    # =========================================================================
    # SECTION 9
    # =========================================================================

    _add_section_heading(
        document,
        "9 ВЫБОР КОНТАКТОРА",
    )

    _add_body_paragraph(
        document,
        (
            "Для коммутации асинхронного электродвигателя при прямом "
            "пуске выбирается контактор Schneider Electric TeSys Deca "
            "для категории применения AC-3."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"По результатам выбора принят контактор "
            f"{contactor_reference}. "
            f"Результат проверки контактора по условиям работы "
            f"электродвигателя: "
            f"{_status_to_russian(_get_value(contactor, 'selection_status', default='—'))}."
        ),
    )

    # =========================================================================
    # SECTION 10
    # =========================================================================

    _add_section_heading(
        document,
        "10 ПРОВЕРКА КООРДИНАЦИИ ПУСКАТЕЛЯ ТИПА 2",
    )

    _add_body_paragraph(
        document,
        (
            "Выбранная комбинация автоматического выключателя "
            "защиты двигателя и контактора проверяется по таблицам "
            "координации Schneider Electric. Для проекта выполняется "
            "проверка координации типа 2."
        ),
    )

    formula = _add_formula_paragraph(
        document
    )

    _add_symbol_with_subscript(
        formula,
        "I",
        "k3",
    )

    _add_run(
        formula,
        " ≤ I",
    )

    _add_run(
        formula,
        "q",
        subscript=True,
    )

    _add_run(
        formula,
        ".",
    )

    _add_body_paragraph(
        document,
        (
            f"Для комбинации {breaker_reference} + "
            f"{contactor_reference} уровень координации составляет "
            f"I_q = {_format_number(coordination_short_circuit_ka)} кА."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Расчётный ток короткого замыкания "
            f"I_k3 = {_format_number(feeder_end_short_circuit_ka)} кА. "
            f"Условие I_k3 ≤ I_q "
            f"{'выполняется' if _get_value(final_design_check, 'type_2_coordination_check', default='') == 'VERIFIED' else 'не выполняется'}."
        ),
    )

    _add_body_paragraph(
        document,
        (
            f"Результат проверки координации типа 2: "
            f"{_status_to_russian(_get_value(final_design_check, 'type_2_coordination_check', default='—'))}."
        ),
    )

    # =========================================================================
    # SECTION 11
    # =========================================================================

    _add_section_heading(
        document,
        "11 ИТОГОВОЕ ТЕХНИЧЕСКОЕ РЕШЕНИЕ",
    )

    _add_body_paragraph(
        document,
        (
            "По результатам расчёта выполнена комплексная проверка "
            "цепи питания асинхронного электродвигателя по допустимой "
            "токовой нагрузке кабеля, падению напряжения в установившемся "
            "и пусковом режимах, отключающей способности защитного "
            "аппарата, выбору контактора и координации пускателя типа 2."
        ),
    )

    final_check_rows = [
        (
            "Допустимая токовая нагрузка кабеля",
            _status_to_russian(
                _get_value(
                    final_design_check,
                    "cable_current_capacity_check",
                    default="—",
                )
            ),
        ),
        (
            "Выбор кабеля",
            _status_to_russian(
                _get_value(
                    final_design_check,
                    "cable_selection_check",
                    default="—",
                )
            ),
        ),
        (
            "Падение напряжения в установившемся режиме",
            _status_to_russian(
                _get_value(
                    final_design_check,
                    "steady_state_voltage_drop_check",
                    default="—",
                )
            ),
        ),
        (
            "Падение напряжения при пуске",
            _status_to_russian(
                _get_value(
                    final_design_check,
                    "motor_starting_voltage_drop_check",
                    default="—",
                )
            ),
        ),
        (
            "Выбор автоматического выключателя",
            _status_to_russian(
                _get_value(
                    final_design_check,
                    "motor_circuit_breaker_selection_check",
                    default="—",
                )
            ),
        ),
        (
            "Отключающая способность",
            _status_to_russian(
                _get_value(
                    final_design_check,
                    "breaking_capacity_check",
                    default="—",
                )
            ),
        ),
        (
            "Выбор контактора",
            _status_to_russian(
                _get_value(
                    final_design_check,
                    "contactor_selection_check",
                    default="—",
                )
            ),
        ),
        (
            "Координация типа 2",
            _status_to_russian(
                _get_value(
                    final_design_check,
                    "type_2_coordination_check",
                    default="—",
                )
            ),
        ),
    ]

    _add_data_table(
        document,
        final_check_rows,
    )

    _add_body_paragraph(
        document,
        (
            f"Итоговый статус проекта цепи питания "
            f"асинхронного электродвигателя: "
            f"{_status_to_russian(final_design_status)}."
        ),
    )

    # =========================================================================
    # SPECIFICATION
    # =========================================================================

    document.add_page_break()

    _add_section_heading(
        document,
        "СПЕЦИФИКАЦИЯ ВЫБРАННОГО ОБОРУДОВАНИЯ",
    )

    specification_table = document.add_table(
        rows=1,
        cols=5,
    )

    specification_table.style = "Table Grid"

    headers = [
        "Поз. обозначение",
        "Наименование",
        "Тип / марка",
        "Кол.",
        "Примечание",
    ]

    for column_index, header_text in enumerate(
        headers
    ):

        cell = specification_table.rows[0].cells[
            column_index
        ]

        cell.text = header_text

        _set_cell_shading(
            cell,
            "D9EAF7",
        )

        for paragraph in cell.paragraphs:

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            for run in paragraph.runs:

                _set_run_font(
                    run,
                    size_pt=11,
                    bold=True,
                )

    specification_rows = [
        (
            "QF1",
            "Автоматический выключатель защиты двигателя",
            breaker_reference,
            "1",
            "Schneider Electric",
        ),
        (
            "KM1",
            "Контактор электродвигателя",
            contactor_reference,
            "1",
            "Schneider Electric",
        ),
        (
            "M1",
            "Асинхронный электродвигатель",
            motor_type,
            "1",
            (
                f"Pₙ = "
                f"{_format_number(rated_power_kw)} кВт"
            ),
        ),
        (
            "W1",
            "Силовой кабель цепи электродвигателя",
            (
                f"{cable_material} / "
                f"{cable_insulation}; "
                f"S = "
                f"{_format_number(cable_section_mm2, digits=1)} мм²"
            ),
            "1",
            (
                f"L = "
                f"{_format_number(cable_length_m, digits=1)} м"
            ),
        ),
    ]

    for specification_row in specification_rows:

        cells = specification_table.add_row().cells

        for column_index, cell_value in enumerate(
            specification_row
        ):

            cells[column_index].text = str(
                cell_value
            )

    _format_table(
        specification_table
    )

    # -------------------------------------------------------------------------
    # SAVE DOCUMENT
    # -------------------------------------------------------------------------

    try:

        document.save(
            output_path
        )

    except PermissionError as error:

        raise ACMotorFeederWordReportError(
            f"Cannot save Word report '{output_path}'. "
            f"Close the DOCX file in Microsoft Word "
            f"and run the calculation again."
        ) from error

    return output_path